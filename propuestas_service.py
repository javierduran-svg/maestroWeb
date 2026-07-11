"""Servicios de propuestas comerciales: numeración, plantillas y exportación."""
from __future__ import annotations

import io
import os
from datetime import date

from fpdf import FPDF

from models import Propuesta

SERVICIOS_PROPUESTA = [
    'CES',
    'CEV+RT',
    'CVS',
    'CES Evaluadora',
    'Consultoría',
    'Proyectos Concesiones hospitalarias',
    'Medición Huella de Carbono',
    'Eficiencia energética',
    'PGSEE',
    'Simulación',
    'Seguimiento en obra',
]

MESES_ES = (
    '', 'Enero', 'Febrero', 'Marzo', 'Abril', 'Mayo', 'Junio',
    'Julio', 'Agosto', 'Septiembre', 'Octubre', 'Noviembre', 'Diciembre',
)

# Tarifas UF por unidad según superficie (CEV+RT)
TARIFAS_CEV_RT = [
    {'label': 'Casas 140 m²', 'm2': 140, 'uf_unidad': 7},
    {'label': 'Casas 200 m²', 'm2': 200, 'uf_unidad': 10},
    {'label': 'Casas 320 m²', 'm2': 320, 'uf_unidad': 16},
]

ETAPAS_CEV_RT = [
    {'codigo': '1.1', 'nombre': 'Informe cumplimiento RT [DOM]'},
    {'codigo': '2.1', 'nombre': 'Pre Calificación'},
    {'codigo': '2.2', 'nombre': 'Calificación'},
]

TEMPLATE_CEV_RT = """<div class="prop-doc">
<div class="prop-doc-header">
  <div class="prop-doc-header-text">
    <h1 class="prop-doc-titulo">Calificación energética de viviendas CEV + Verificación Reglamentación térmica.</h1>
    <h2 class="prop-doc-subtitulo">{{PROYECTO}}</h2>
  </div>
  <div class="prop-doc-logo-wrap">{{LOGO}}</div>
</div>
<table class="prop-doc-meta">
  <tr><th>Cliente:</th><td>{{CLIENTE}}</td></tr>
  <tr><th>Presentada por:</th><td>{{PRESENTADO_POR}}</td></tr>
  <tr><th>Fecha:</th><td>{{FECHA}}</td></tr>
  <tr><th>ID Propuesta:</th><td>P{{NUMERO}}</td></tr>
</table>

<h3 class="prop-doc-seccion">Introducción</h3>
<p>La presente Propuesta Técnica se desarrolla para el proyecto <strong>{{PROYECTO}}</strong>, y tiene por objetivo dar cumplimiento a los requerimientos normativos vigentes en materia de Reglamentación Térmica y Calificación Energética de Viviendas (CEV).</p>
<p>El encargo considera la elaboración de los informes técnicos exigidos por la Dirección de Obras Municipales (DOM) para el ingreso y aprobación de modificaciones de proyecto, así como la evaluación energética integral del conjunto habitacional.</p>
<p>Adicionalmente, la propuesta incluye la Precalificación y Calificación Energética CEV de cada una de las viviendas del condominio, correspondiente a <strong>{{UNIDADES_DESCRIPCION}}</strong>, a evaluar con el objetivo de optimizar su desempeño energético y alcanzar la mejor calificación posible dentro del marco normativo vigente.</p>

<h3 class="prop-doc-seccion">Propuesta Técnica</h3>
<p>La presente propuesta técnica se estructura en dos componentes principales, orientados a verificar el cumplimiento normativo y evaluar el desempeño energético del proyecto en etapa de diseño.</p>

<h4>1. Cumplimiento de la Reglamentación Térmica</h4>
<p>Se elaborará un Informe de Cumplimiento de Reglamentación Térmica válido para presentación ante la Dirección de Obras Municipales (DOM), en el cual se verificará el cumplimiento del Artículo 4.1.10 de la OGUC, aplicable a edificaciones de uso residencial.</p>
<p><strong>A. Desempeño térmico de la envolvente</strong> — Se verificará el cumplimiento de transmitancia térmica (U) o resistencia térmica (Rt) de techumbres, muros, pisos, puertas y ventanas, con memoria de cálculo detallada.</p>
<p><strong>B. Ausencia de riesgo de condensación</strong> — Memoria de cálculo con método de Glaser para muros, cubiertas y pisos ventilados.</p>
<p><strong>C. Permeabilidad al aire e infiltraciones</strong> — Revisión de carpinterías, sellados y barreras de vapor. Blower door no incluida (servicio adicional).</p>
<p><strong>D. Ventilación mínima</strong> — Diseño conceptual de ventilación conforme a NCh 3308.</p>

<h4>2. Calificación energética de viviendas CEV</h4>
<p><strong>2.1. Análisis Preliminar CEV</strong> — Simulación temprana y propuestas de mejora.</p>
<p><strong>2.2. Precalificación CEV</strong> — Simulación PBDT MINVU con permiso de edificación aprobado.</p>
<p><strong>2.3. Calificación energética</strong> — Visita de obra obligatoria en construcción.</p>
<p><strong>2.4. Calificación CEV</strong> — Tramitación final tras Recepción Final.</p>

<h3 class="prop-doc-seccion">Honorarios Profesionales</h3>
<p>Para definir el monto de los honorarios profesionales se asume que se contratan los 2 servicios descritos en la propuesta:</p>
{{HONORARIOS_TABLA}}

<h4>Forma de pago</h4>
{{PAGO_TABLA}}
<p class="prop-doc-total"><strong>TOTAL: UF {{TOTAL_UF}}</strong></p>

<div class="prop-doc-firma">
  <p><strong>{{PRESENTADO_POR}}</strong></p>
  <p>Arquitecto PUC | Master en Medio Ambiente y Arquitectura Bioclimática U. Politécnica de Madrid |<br>
  LEED AP | Asesor CES | Calificador Energético CEV.<br>
  B-green Chile</p>
</div>
<div class="prop-doc-empresa">
  <p><strong>Información de la Empresa</strong></p>
  <p>Nombre: B-green Chile Ltda.<br>
  Rut.: 77.748.415-k<br>
  Dirección: Obispo Donoso 5 Oficina 62. Providencia.</p>
</div>
</div>"""

TEMPLATES_POR_SERVICIO = {
    'CEV+RT': TEMPLATE_CEV_RT,
}


def siguiente_numero_propuesta(empresa_id: int) -> int:
    max_num = (
        Propuesta.query.filter_by(empresa_id=empresa_id)
        .with_entities(Propuesta.numero)
        .order_by(Propuesta.numero.desc())
        .limit(1)
        .scalar()
    )
    return (max_num or 0) + 1


def _texto_seguro(texto) -> str:
    if texto is None:
        return ''
    s = str(texto)
    for a, b in (
        ('á', 'a'), ('é', 'e'), ('í', 'i'), ('ó', 'o'), ('ú', 'u'),
        ('Á', 'A'), ('É', 'E'), ('Í', 'I'), ('Ó', 'O'), ('Ú', 'U'),
        ('ñ', 'n'), ('Ñ', 'N'), ('ü', 'u'), ('Ü', 'U'),
    ):
        s = s.replace(a, b)
    return s


def _fmt_fecha_larga(fecha_str: str | None) -> str:
    if not fecha_str:
        hoy = date.today()
        return f'{hoy.day:02d} de {MESES_ES[hoy.month]} de {hoy.year}'
    try:
        partes = str(fecha_str)[:10].split('-')
        if len(partes) == 3:
            anio, mes, dia = int(partes[0]), int(partes[1]), int(partes[2])
            if 1 <= mes <= 12:
                return f'{dia:02d} de {MESES_ES[mes]} de {anio}'
    except (ValueError, IndexError):
        pass
    return str(fecha_str)


class PropuestaPDF(FPDF):
    def __init__(self):
        super().__init__()
        self.set_margins(20, 20, 20)
        self.set_auto_page_break(auto=True, margin=20)

    def header(self):
        pass


def _html_a_texto(html: str) -> str:
    import re
    import html as html_mod

    texto = re.sub(r'<br\s*/?>', '\n', html, flags=re.I)
    texto = re.sub(r'</p>', '\n\n', texto, flags=re.I)
    texto = re.sub(r'</h[1-6]>', '\n\n', texto, flags=re.I)
    texto = re.sub(r'</tr>', '\n', texto, flags=re.I)
    texto = re.sub(r'</t[dh]>', '\t', texto, flags=re.I)
    texto = re.sub(r'<[^>]+>', '', texto)
    texto = html_mod.unescape(texto)
    texto = re.sub(r'\n{3,}', '\n\n', texto)
    return texto.strip()


def generar_pdf_propuesta(titulo: str, contenido: str, logo_path: str | None = None) -> bytes:
    pdf = PropuestaPDF()
    pdf.add_page()
    y_ini = pdf.get_y()
    if logo_path and os.path.isfile(logo_path):
        try:
            pdf.image(logo_path, x=150, y=y_ini, h=18)
        except Exception:
            pass
    pdf.set_font('Helvetica', 'B', 14)
    pdf.multi_cell(0, 8, _texto_seguro(titulo))
    pdf.ln(4)
    pdf.set_font('Helvetica', '', 10)
    texto = _html_a_texto(contenido) if '<' in contenido else contenido
    for linea in texto.split('\n'):
        if linea.strip():
            pdf.multi_cell(0, 5, _texto_seguro(linea))
        else:
            pdf.ln(3)
    return bytes(pdf.output())


def generar_docx_propuesta(
    titulo: str, contenido: str, logo_path: str | None = None,
) -> tuple[bytes, str]:
    """Retorna (bytes, extension: 'docx' | 'doc')."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches

        doc = Document()
        if logo_path and os.path.isfile(logo_path):
            try:
                doc.add_picture(logo_path, width=Inches(1.6))
            except Exception:
                pass
        titulo_p = doc.add_heading(titulo, level=1)
        titulo_p.runs[0].font.size = Pt(16)
        texto = _html_a_texto(contenido) if '<' in contenido else contenido
        for linea in texto.split('\n'):
            p = doc.add_paragraph(linea)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.size = Pt(10)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue(), 'docx'
    except ImportError:
        return _generar_doc_html(titulo, contenido), 'doc'


def _generar_doc_html(titulo: str, contenido: str) -> bytes:
    import html
    body = html.escape(contenido).replace('\n', '<br>')
    titulo_esc = html.escape(titulo)
    doc_html = (
        f'<html xmlns:o="urn:schemas-microsoft-com:office:office" '
        f'xmlns:w="urn:schemas-microsoft-com:office:word">'
        f'<head><meta charset="utf-8"><title>{titulo_esc}</title></head>'
        f'<body><h1>{titulo_esc}</h1><div style="font-family:Calibri;font-size:11pt">{body}</div></body></html>'
    )
    return doc_html.encode('utf-8')


def get_config_calculadora(servicio: str) -> dict | None:
    if servicio == 'CEV+RT':
        return {
            'tarifas': TARIFAS_CEV_RT,
            'etapas': ETAPAS_CEV_RT,
            'template': TEMPLATE_CEV_RT,
            'format': 'html',
        }
    return None
